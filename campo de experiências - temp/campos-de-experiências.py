import json
import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def carregar_e_limpar_dados(caminho_json: Path) -> list:
    """
    Carrega os dados de um arquivo JSON, realiza o parse resiliente caso haja 
    múltiplos blocos concatenados, limpa citações e ordena cronologicamente.
    Salva uma versão limpa de volta no arquivo.
    """
    # Verifica se o arquivo existe
    if not caminho_json.exists():
        print(f"[ERRO] O arquivo 'dados.json' não foi encontrado em:\n{caminho_json}")
        return []

    # Lê o conteúdo bruto do arquivo para processamento resiliente
    try:
        conteudo = caminho_json.read_text(encoding='utf-8')
    except Exception as e:
        print(f"[ERRO] Não foi possível ler o arquivo 'dados.json': {e}")
        return []

    # Parser resiliente para extrair múltiplos blocos JSON colados ou malformados
    dados = []
    decoder = json.JSONDecoder()
    pos = 0
    tamanho = len(conteudo)

    while pos < tamanho:
        # Pula caracteres que não iniciam uma estrutura JSON (lista ou objeto)
        while pos < tamanho and conteudo[pos] not in ('[', '{'):
            pos += 1
        if pos >= tamanho:
            break
        try:
            objeto, idx = decoder.raw_decode(conteudo[pos:])
            if isinstance(objeto, list):
                dados.extend(objeto)
            elif isinstance(objeto, dict):
                dados.append(objeto)
            pos += idx
        except json.JSONDecodeError:
            pos += 1

    if not dados:
        print("[ERRO] Nenhum dado JSON válido pôde ser extraído do arquivo 'dados.json'.")
        return []

    # Limpeza de citações e textos indesejados (ex: [cite: 13])
    for item in dados:
        for campo in ['data', 'campo', 'experiencia']:
            if campo in item and isinstance(item[campo], str):
                # Remove o padrão [cite: ...] e espaços extras nas pontas
                item[campo] = re.sub(r'\s*\[cite:[^\]]+\]', '', item[campo]).strip()

    # Reordenação cronológica por data
    try:
        dados.sort(key=lambda x: datetime.strptime(x.get('data', '01/01/70'), '%d/%m/%y'))
    except Exception:
        # Se falhar na ordenação (por datas fora do padrão), tenta ordenar como strings simples
        try:
            dados.sort(key=lambda x: x.get('data', ''))
        except Exception:
            pass

    # Reescreve o arquivo dados.json formatado como uma lista JSON única e válida
    try:
        caminho_json.write_text(json.dumps(dados, ensure_ascii=False, indent=2), encoding='utf-8')
        print("[INFO] Arquivo 'dados.json' reorganizado, limpo de citações e salvo com sucesso!")
    except Exception as e:
        print(f"[AVISO] Não foi possível reescrever o arquivo 'dados.json' limpo: {e}")

    return dados


def gerar_documento_word(dados: list) -> Document:
    """
    Cria um documento Word formatado com os dados das experiências em uma tabela.
    """
    doc = Document()

    # CONFIGURAÇÃO DE PÁGINA: PAISAGEM
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # CONFIGURAÇÃO DE FONTE: ARIAL
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # TABELA
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # CABEÇALHO
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Data'
    hdr_cells[1].text = 'Campos de experiências'
    hdr_cells[2].text = 'Experiências desenvolvidas'

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.bold = True
            run.font.name = 'Arial'

    # Ajuste de largura
    for cell in table.columns[0].cells:
        cell.width = Cm(2.5)
    for cell in table.columns[1].cells:
        cell.width = Cm(7.0)
    for cell in table.columns[2].cells:
        cell.width = Cm(17.2)

    # PREENCHIMENTO
    for item in dados:
        row_cells = table.add_row().cells

        row_cells[0].text = item.get('data', '-')
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_cells[1].text = item.get('campo', '-')
        row_cells[2].text = item.get('experiencia', '-')

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles['Normal']

    return doc


def salvar_documento(doc: Document, diretorio_base: Path):
    """
    Cria a pasta de destino caso necessário e salva o documento Word.
    """
    pasta_destino = diretorio_base / 'Campos de experiência'

    # Cria a pasta se ela não existir
    if not pasta_destino.exists():
        pasta_destino.mkdir(parents=True, exist_ok=True)
        print(f"[PASTA] Pasta criada: {pasta_destino}")
    else:
        print(f"[PASTA] Usando pasta existente: {pasta_destino}")

    # Gera o nome do arquivo com a data atual
    data_hoje = datetime.now().strftime("%d-%m-%Y")
    nome_arquivo = f'Campos_de_experiências_{data_hoje}.docx'
    caminho_completo = pasta_destino / nome_arquivo

    try:
        doc.save(str(caminho_completo))
        print(f"[SUCESSO] Arquivo salvo localmente em:\n{caminho_completo}")
    except Exception as e:
        print(f"[ERRO] Não foi possível salvar o arquivo Word: {e}")


if __name__ == "__main__":
    # Resolve o diretório onde o script está localizado
    diretorio_script = Path(__file__).parent if '__file__' in locals() else Path.cwd()
    caminho_json = diretorio_script / 'dados.json'

    # 1. Carrega, limpa e formata os dados do JSON
    dados_aulas = carregar_e_limpar_dados(caminho_json)

    # 2. Se houver dados válidos, gera o documento e salva
    if dados_aulas:
        documento = gerar_documento_word(dados_aulas)
        salvar_documento(documento, diretorio_script)
